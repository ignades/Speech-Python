
#Python speech to text and text to speech Ignacio Olivieri upgrade software of another programmer v1.0 beta

import locale
import speech_recognition as sr
import os
import argparse
import sys
import logging
import docx
from option_name import *
from inquirer_form_loop import *
from write_doc import *
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Initialize the recognizer

r = sr.Recognizer()

# speech
def record_text():
    #loop in case of error
    while(1):
        try:
            # use the microphone as source for input.
            with sr.Microphone() as source2:
                #Prepare
                r.adjust_for_ambient_noise(source2)
                #listens for the user's input
                audio2 = r.listen(source2,phrase_time_limit=10)
                # Using google to recognize audio
                MyText = r.recognize_google(audio2,language="it-IT")
                #return MyText = MyText.lower()
                print(MyText)
                #os.system(MyText)
                return MyText

        except sr.RequestError as e:
            print("Could not request results; {0}".format(e))

        except sr.UnknownValueError:
            print("unknown error occurred")
    return

def scrivi(doc):
    print("Scrivendo..")
    text2 = record_text()
    # Add another paragraph
    p = doc.add_paragraph()
    # Add a run and format it
    run = p.add_run(text2)
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)

    #print("Scrivendo ancora...")
    #f=open("output_text.txt","a")
    #f.write(text2)
    #f.write("\n")
    #f.close()

    return

def selection():
    #os.system("gnome-terminal -- 'bash -c \""+command+""'")
    text = record_text()
    select = os.system("python3 option.py")
    print(select)
    return select

def crea_doc():
    #text = record_text()
    print("CREA DOC ...")
    doc = docx.Document()
    assegna_nome = give_name()
    # Add a paragraph to the document
    p = doc.add_paragraph()
    # Add some formatting to the paragraph
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add a run to the paragraph
    run = p.add_run(assegna_nome+"\n")

    # Add some formatting to the run
    run.bold = True
    run.italic = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(16)
    doc.save("./Documents/"+assegna_nome+".docx")
    return


def exit_program():
    print("Exiting the program...")
    sys.exit(0)

while(1):
    text = record_text()

    if text == "Special":
        select = selection()
    elif text == "crea documento":
        print("Crea..")
        name_doc = crea_doc()
    elif text=="scrivi":
        doc_to_edit = create_inquirer()
        doc_edit = doc_to_edit.replace("./Documents/", "")
        testo = record_text()
        write_documento(doc_edit,testo)
    elif text == "esci":
        exit_program()
    else:
        select = "No Text"
