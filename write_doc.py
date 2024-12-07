import docx
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

def write_documento(name_doc,testo):
    try:
        doc = Document(name_doc)
        print(name_doc)

        if "punto a capo" in testo:
             xx = testo.replace("punto a capo",".")
             testo = xx

        # Add a run to the paragraph
        if "virgola" in testo:
             yy = testo.replace("virgola",",")
             testo = yy

        print(testo)
        # Add another paragraph

        # Add a run and format it
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p2.add_run(testo)
        run.italic = True
        run.font.name = 'Arial'

        run.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

        run.font.size = docx.shared.Pt(12)


    except():
        print("Could not request results")



    doc.save("./Documents/"+name_doc)
    print("Documento! ok")



#write_doc("pippo.docx","prova primo paragrafo virgola altro testo punto a capo")
