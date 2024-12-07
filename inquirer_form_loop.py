import docx
import inquirer
import array
import glob
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from option_name import *

def crea_doc():
    #text = record_text()
    print("CREA DOC ...")
    doc = docx.Document()
    assegna_nome = give_name()
    # Add a paragraph to the document
    p = doc.add_paragraph()
    # Add some formatting to the paragraph
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add a run to the paragraph
    run = p.add_run(assegna_nome+"\n")

    # Add some formatting to the run
    run.bold = True
    run.italic = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(16)
    doc.save("./Documents/"+assegna_nome+".docx")
    return

def create_inquirer():
    docs =  glob.glob('./Doduments/*.docx')
    l = len(docs)
    #return print(l)
    all_choices = []

    if(l == 0):
        message= "Crea un nuovo Documento!"
        print(message)
        crea_doc()
        return  message
    else:

        for idx, x in enumerate(docs):
            all_choices.insert(idx, x)

        questions = [
          inquirer.List('size',
                        message="Choice a document to edit?",
                        choices=all_choices,
                    ),
        ]

        answers = inquirer.prompt(questions)
        print(answers['size'])
        return  answers['size']


#create_inquirer()
